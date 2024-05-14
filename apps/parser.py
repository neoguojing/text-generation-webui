
import json
from pathlib import Path
import re
from typing import Any, List, Mapping, Optional,Dict,Union,Tuple
from pydantic import  Field, BaseModel,validator
from langchain.schema.agent import AgentAction, AgentFinish
from langchain.agents.agent import AgentOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader,TextLoader

import os
import time
# 获取当前脚本所在的目录路径
current_dir = os.path.dirname(os.path.abspath(__file__))

# 将当前package的父目录作为顶层package的路径
top_package_path = os.path.abspath(os.path.join(current_dir, ".."))

class QAItem(BaseModel):
    question: str = Field(description="question")
    answer: str = Field(description="answer")

    
class QAPackage(BaseModel):
    data: List[QAItem] = Field(..., description="问题答案列表")

    def merge_data(self, other: 'QAPackage'):
        self.data.extend(other.data)

    def length(self):
        return len(self.data)
    
    def dump(self):
        dict_obj = self.dict()
        return json.dumps(dict_obj["data"], ensure_ascii=False, indent=4)
    
    def load(self, path: str) -> Optional[dict]:
        try:
            with open(path) as f:
                json_data = json.load(f)
                qa_items = [QAItem(**item) for item in json_data]
                self.data=qa_items
        except FileNotFoundError:
            return None
        
    def toQwen(self,source:str):
        qwen_item = QwenItem(id=source,conversations=[])
        for qa in self.data:
            if qa.answer == "":
                continue
            data_input = {"from": "user", "value": qa.question}
            q = QwenConversationItem(**data_input)
            qwen_item.conversations.append(q)
            data_input = {"from": "assistant", "value": qa.answer}
            a = QwenConversationItem(**data_input)
            qwen_item.conversations.append(a)

        return qwen_item.dump()
    
    def toLLama(self,path,source:str):
        with open(path+source+"_llama"+".json", 'w', encoding='utf-8') as f:
            for qa in self.data:
                if qa.answer == "":
                    continue
                data_input = {"instruction": "", "input": qa.question,"output":qa.answer}
                q = LlamaItem(**data_input)
                json_str = q.dump()
                f.write(json_str+"\n")


        

class LlamaItem(BaseModel):
    instruction: str = Field(...,  description="指令")
    input: str = Field(..., description="输入")
    output: str = Field(..., description="输出")

    def dump(self):
        dict_obj = self.dict()
        return json.dumps(dict_obj, ensure_ascii=False, indent=2)


class QwenConversationItem(BaseModel):
    from_: str = Field(..., alias="from", description="发送方")
    value: str = Field(..., description="消息内容")


class QwenItem(BaseModel):
    id: str = Field(..., description="标识")
    conversations: List[QwenConversationItem] = Field(..., description="对话列表")

    def merge_data(self, other: 'QwenItem'):
        self.conversations.extend(other.conversations)

    def length(self):
        return len(self.conversations)
    
    def dump(self):
        dict_obj = self.dict()
        return json.dumps(dict_obj, ensure_ascii=False, indent=2)
    
    def toLLama(self,path,source:str):
        with open(path+source+"_llama"+".json", 'w', encoding='utf-8') as f:
            for i in range(0, len(self.conversations), 2):
                data_input = {"instruction": "", "input": self.conversations[i].value,"output":self.conversations[i+1].value}
                q = LlamaItem(**data_input)
                json_str = q.dump()
                f.write(json_str+"\n")

class JsonOutputParser(AgentOutputParser):
    pattern = re.compile(r"```(?:json)?\n(.*?)```", re.DOTALL)
    qaList: QAPackage = QAPackage(data=[])

    class Config:
        arbitrary_types_allowed = True

    def parse(self, llm_output: str) -> Union[AgentAction, AgentFinish]:
        print("llm_output--------",llm_output)
        data = None
        # Check if the output contains valid JSON
        try:
            action_match = self.pattern.search(llm_output)
            if action_match is not None:
                response = action_match.group(1).strip()
            else:
                response = llm_output
            response = json.loads(response, strict=False)
            data = response
        except json.JSONDecodeError:
            print("***********Invalid JSON in LLM output")
        
        # Parse the JSON into a dictionary
        print("data-----",data)
        output = {}
        if isinstance(data, dict):
            output = data
            tmp = QAPackage(data=data["data"])
            print("step qa num:",tmp.length())
            self.qaList.merge_data(tmp)
            print("total:",self.qaList.length())
        elif isinstance(data, list):
            output = {"data": data}
            tmp = QAPackage(data=data)
            print("step qa num:",tmp.length())
            self.qaList.merge_data(tmp)
            print("total:",self.qaList.length())
        return AgentFinish(return_values=output,log=llm_output)
    
    def dump(self, path: str):
        print("final:",self.qaList.length())
        with open(path+".json", 'w', encoding='utf-8') as f:
            package_json = self.qaList.dump()
            f.write(package_json)
        with open(path+".qwen", 'w', encoding='utf-8') as f:
            package_json = self.qaList.toQwen(path)
            f.write(package_json)

        self.qaList = QAPackage(data=[])
            
    def load(self, path: str) -> Optional[dict]:
        try:
            with open(path) as f:
                return json.load(f)
        except FileNotFoundError:
            return None


def data_generate_chain(data_dir: str,glob: str = "**/*.txt",model_type: str="tongyi",):
    from langchain.output_parsers import PydanticOutputParser
    from apps.model_factory import ModelFactory
    from apps.parser import JsonOutputParser,QAPackage,QAItem
    from apps.prompt import PromptFactory
    from typing import Any, List

    loader = DirectoryLoader(data_dir, glob=glob,loader_cls=TextLoader)
    # loader = TextLoader("./doc.txt")
    docs = loader.load()

    # llm = ModelFactory().get_model("openai")
    # llm = ModelFactory().get_model("claude")
    # llm = ModelFactory().get_model("qwen")
    # llm = ModelFactory().get_model("qianfan")
    llm = ModelFactory.get_model(model_type)
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, chunk_overlap=0
    )
    
    qaParser = PydanticOutputParser(pydantic_object=QAPackage)

    prompt = PromptFactory.caibao_analyse_prompt(qaParser.get_format_instructions())
    
    texts = []
    for doc in docs:
        text = doc.page_content
        print(doc.metadata)
        jsonParser = JsonOutputParser()
        chain = prompt | llm | jsonParser 

        texts += text_splitter.create_documents([text])
        for text in texts:
            print(text)
            try:
                answer = chain.invoke({"text": text,"format_instructions":qaParser.get_format_instructions()})
            except Exception as e:
                print("",str(e))
                continue
            
            print(f"Output: {answer}")
            time.sleep(1)
        
        jsonParser.dump(os.path.splitext(doc.metadata["source"])[0])

import pdb
class QwenAgentOutputParser(AgentOutputParser):

    def parse(self, llm_output: str) -> Union[AgentAction, AgentFinish]:
        # pdb.set_trace()
        print("llm_output---------:\n",llm_output)
        # Check if agent should finish
        if "Final Answer:" in llm_output:
            return AgentFinish(
                # Return values is generally always a dictionary with a single `output` key
                # It is not recommended to try anything else at the moment :)
                return_values={"output": llm_output.split("Final Answer:")[-1].strip()},
                log=llm_output,
            )
        
        # Parse out the action and action input
        # regex = r"Action\s*\d*\s*:(.*?)\nAction\s*\d*\s*Input\s*\d*\s*:[\s]*(.*)"
        regex = r"Action\s*\d*\s*:[\s]*(.*?)[\s]*Action\s*\d*\s*Input\s*\d*\s*:[\s]*(.*)"
        match = re.search(regex, llm_output, re.DOTALL)

        if not match:
            # raise ValueError(f"Could not parse LLM output: `{llm_output}`")
            return AgentFinish(
                    # Return values is generally always a dictionary with a single `output` key
                    # It is not recommended to try anything else at the moment :)
                    return_values={"output": llm_output.strip()},
                    log=llm_output,
                )
        # print("llm_output:",llm_output)
        action = match.group(1).strip()
        print("action:---------",action)
        action_input = match.group(2)
        print("action_input--------:",action)
        tool_input = action_input.strip('\nObservation: ').strip(" ").strip('"')
        print(f"tool_input--------:{tool_input}")
        # Return the action and action input
        return AgentAction(tool=action, tool_input=tool_input, log=llm_output)

def parse_latest_plugin_call(text: str) -> Tuple[str, str]:
    i = text.rfind('\nAction:')
    j = text.rfind('\nAction Input:')
    k = text.rfind('\nObservation:')
    if 0 <= i < j:  # If the text has `Action` and `Action input`,
        if k < j:  # but does not contain `Observation`,
            # then it is likely that `Observation` is ommited by the LLM,
            # because the output text may have discarded the stop word.
            text = text.rstrip() + '\nObservation:'  # Add it back.
            k = text.rfind('\nObservation:')
    if 0 <= i < j < k:
        plugin_name = text[i + len('\nAction:'):j].strip()
        plugin_args = text[j + len('\nAction Input:'):k].strip()
        return plugin_name, plugin_args
    return '', ''

def use_api(tools, response):
    use_toolname, action_input = parse_latest_plugin_call(response)
    if use_toolname == "":
        return "no tool founds"

    used_tool_meta = list(filter(lambda x: x["name_for_model"] == use_toolname, tools))
    if len(used_tool_meta) == 0:
        return "no tool founds"
    
    api_output = used_tool_meta[0]["tool_api"](action_input)
    return api_output

def parse_table_to_string(table):
    table_string = ""
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n','')
            cell_string = f"{cell_text} | "
            table_string += cell_string
        table_string += "\n"
    return table_string

def docx_parser(file_path: str,dst_path: str):
    from docx import Document
    import pandas as pd
    document = Document(file_path)
    with open(dst_path, 'w', encoding='utf-8') as f:
        i,j = 0,0
        for element in document.element.body:
            if element.tag.endswith('p'):  # Paragraph
                f.write(document.paragraphs[j].text.strip()+"\n")
                j+=1
            elif element.tag.endswith('tbl'):  # Table
                # pdb.set_trace()
                table_str = parse_table_to_string(document.tables[i])
                f.write(table_str)
                i+=1

# if __name__ == '__main__':
    # qas = QAPackage(data=[])
    # qas.load("./ir2023_ashare.json")
    # qas.toLLama(".","ir2023_ashare")
    # output = qas.toQwen("ir2023_ashare.qw")
    # with open("ir2023_ashare.qw", 'w', encoding='utf-8') as f:
    #     f.write(output)

    # with open("../dataset/chat/ir2023_ashare.qwen") as f:
    #     data = f.read()
    #     qw = QwenItem.parse_raw(data) 
    #     qw.toLLama("./","ir2023_ashare")
    
    # from langchain.document_loaders import PyPDFLoader
    # # from langchain.document_loaders import UnstructuredPDFLoader
    # from langchain.document_loaders import PDFMinerPDFasHTMLLoader
    # from langchain.document_loaders import PDFPlumberLoader
    # loader = PDFPlumberLoader("../dataset/chat/ir2023_ashare.pdf")
    # # pages = loader.load_and_split()
    # text_splitter = RecursiveCharacterTextSplitter(
    #     chunk_size=500, chunk_overlap=50,
    # )
    # pages = loader.load()
    # with open("test.txt", 'w', encoding='utf-8') as f:
    #     for p in pages:
    #         print(p.page_content)
    #         # texts = text_splitter.create_documents([p.page_content])
    #         # for t in texts:
    #         #     print(t)
    #         f.write(p.page_content)

    # docx_parser("../dataset/chat/ir2023_ashare.docx","text.txt")

    # loader = TextLoader("./text.txt")

    # text_splitter = RecursiveCharacterTextSplitter(
    #     chunk_size=500, chunk_overlap=50,
    # )
    # pages = loader.load()
    # with open("splite.txt", 'w', encoding='utf-8') as f:
    #     for p in pages:
    #         print(p.page_content)
    #         texts = text_splitter.create_documents([p.page_content])
    #         for t in texts:
    #             print(t)
    #             f.write(t.page_content)